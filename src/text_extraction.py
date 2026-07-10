# =====================================================================
# BLOCK: TEXT EXTRACTION (src/text_extraction.py)
# ---------------------------------------------------------------------
# WHAT THIS FILE IS:
#   One job only: take a file (PDF / Word / image / text) and return
#   the plain text inside it. Nothing here knows about scoring,
#   candidates, or the user interface.
#
# NOTE: We use PyMuPDF ("fitz") for PDFs. The old code used TWO PDF
#   libraries (PyPDF2 + PyMuPDF) doing the same job; PyPDF2 is also
#   deprecated, so it was removed.
# =====================================================================

import logging
import os

import docx                      # reads .docx Word files
import fitz                      # PyMuPDF — reads PDF files
import pytesseract               # OCR — reads text out of images
from PIL import Image

logger = logging.getLogger(__name__)


def extract_text(path: str) -> str:
    """Return the plain text of a CV file, or '' if it cannot be read.

    Supported: .pdf, .docx, .png, .jpg, .jpeg, .txt
    (Old .doc files are NOT supported — ask candidates for PDF/DOCX,
    or open the file in Word and re-save it as .docx.)
    """
    ext = os.path.splitext(path)[1].lower()
    try:
        if ext == ".pdf":
            return _read_pdf(path)
        if ext == ".docx":
            return _read_docx(path)
        if ext in (".png", ".jpg", ".jpeg"):
            return _read_image(path)
        if ext == ".txt":
            return _read_txt(path)
    except Exception as exc:  # a single corrupt file must never crash the app
        logger.error("فشلت قراءة الملف %s. السبب: %s", path, exc)
    return ""


# --------------------- one small helper per format -------------------

def _read_pdf(path: str) -> str:
    with fitz.open(path) as doc:
        text = "\n".join(page.get_text() for page in doc)
    # Scanned PDFs contain images, not text. If we got (almost) nothing,
    # fall back to OCR on each page so scanned CVs still work.
    if len(text.strip()) >= 40:
        return text
    with fitz.open(path) as doc:
        ocr_parts = []
        for page in doc:
            pix = page.get_pixmap(dpi=200)
            img = Image.frombytes("RGB", (pix.width, pix.height), pix.samples)
            ocr_parts.append(pytesseract.image_to_string(img, lang="eng+ara"))
    return "\n".join(ocr_parts)


def _read_docx(path: str) -> str:
    document = docx.Document(path)
    parts = [para.text for para in document.paragraphs]
    # The old code ignored tables — many CVs keep contact details in a
    # table, so emails/phones were silently lost. Read them too.
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def _read_image(path: str) -> str:
    return pytesseract.image_to_string(Image.open(path), lang="eng+ara")


def _read_txt(path: str) -> str:
    with open(path, "r", encoding="utf-8", errors="ignore") as handle:
        return handle.read()
